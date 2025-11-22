import os
import types
from pathlib import Path

import pytest

from src import processor


def test_extract_text_from_pdf(monkeypatch, tmp_path):
    class FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class FakeReader:
        def __init__(self, file_obj):
            self.pages = [FakePage("Page 1"), FakePage("Page 2")]

    monkeypatch.setattr(processor.pypdf, "PdfReader", FakeReader)

    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    text = processor.extract_text_from_pdf(str(pdf_path))

    assert text == "Page 1Page 2"


def test_extract_text_from_docx(tmp_path):
    docx_path = tmp_path / "sample.docx"
    from docx import Document

    document = Document()
    document.add_paragraph("First line")
    document.add_paragraph("Second line")
    document.save(docx_path)

    text = processor.extract_text_from_docx(str(docx_path))

    assert text == "First line\nSecond line\n"


def test_extract_structured_data_missing_api_key(monkeypatch):
    monkeypatch.delenv("GOOGLE_API_KEY", raising=False)

    result = processor.extract_structured_data("Some resume text")

    assert result is None


def test_extract_structured_data_with_stubs(monkeypatch):
    monkeypatch.setenv("GOOGLE_API_KEY", "test-key")

    class FakeModel:
        def __init__(self, data):
            self._data = data

        def dict(self):
            return self._data

    class FakeParser:
        def __init__(self, pydantic_object):
            self._pydantic_object = pydantic_object

        def get_format_instructions(self):
            return "<format>"

        def parse(self, text):
            return FakeModel({"ai_summary": text, "name": "Parsed"})

    class FakeChain:
        def __init__(self, handler):
            self._handler = handler

        def __or__(self, other):
            def handler(params):
                return other.parse(self._handler(params))

            return FakeChain(handler)

        def invoke(self, params):
            return self._handler(params)

    class FakePrompt:
        def __init__(self, template, input_variables, partial_variables):
            self.template = template

        def __or__(self, other):
            def handler(params):
                return other.generate(params["text"])

            return FakeChain(handler)

    class FakeLLM:
        def __init__(self, model, google_api_key, temperature):
            self.model = model
            self.google_api_key = google_api_key
            self.temperature = temperature

        def generate(self, text):
            return f"structured: {text}"

    monkeypatch.setattr(processor, "PydanticOutputParser", FakeParser)
    monkeypatch.setattr(processor, "PromptTemplate", FakePrompt)
    monkeypatch.setattr(processor, "ChatGoogleGenerativeAI", FakeLLM)

    result = processor.extract_structured_data("Resume text")

    assert result == {"ai_summary": "structured: Resume text", "name": "Parsed"}


def test_process_resumes_skips_existing(monkeypatch, tmp_path):
    resume_dir = tmp_path / "resumes"
    resume_dir.mkdir()
    pdf_path = resume_dir / "existing.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    monkeypatch.setattr(processor, "get_candidate_by_filename", lambda filename: {"id": 1})

    def fake_extract(*args, **kwargs):  # pragma: no cover - should not be called
        raise AssertionError("extract_text_from_pdf should not be called for existing candidates")

    monkeypatch.setattr(processor, "extract_text_from_pdf", fake_extract)

    processor.process_resumes(str(resume_dir))


def test_process_resumes_adds_new_candidate(monkeypatch, tmp_path):
    resume_dir = tmp_path / "resumes"
    resume_dir.mkdir()
    pdf_path = resume_dir / "new.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    captured_candidates = []

    monkeypatch.setattr(processor, "get_candidate_by_filename", lambda filename: None)
    monkeypatch.setattr(processor, "extract_text_from_pdf", lambda filepath: "resume content")
    monkeypatch.setattr(
        processor,
        "extract_structured_data",
        lambda text: {"name": "Test User", "work_experience": []},
    )
    monkeypatch.setattr(processor, "add_candidate", lambda data: captured_candidates.append(data))

    processor.process_resumes(str(resume_dir))

    assert captured_candidates == [
        {"name": "Test User", "work_experience": [], "filename": "new.pdf"}
    ]
